import io
import json
import urllib.parse
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
import google.generativeai as genai
from PIL import Image
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.ensemble import RandomForestClassifier

# ==========================================
# 0. СЛОВАРЬ ПЕРЕВОДОВ ИНТЕРФЕЙСА
# ==========================================
translations = {
    "ru": {
        "page_title": "AI-Помощник Учителя",
        "sidebar_title": "🎓 EduAI Platform",
        "api_subheader": "🔑 Доступ к ИИ",
        "api_help": "Введите ключ один раз для всех инструментов",
        "api_expander": "ℹ️ Как получить API ключ бесплатно?",
        "tools_subheader": "🛠️ Модули системы",
        "footer": "✨ Разработано для преподавателей",
        "menu": [
            "📝 Генератор карточек",
            "📅 AI-Генератор КТП",
            "📋 AI-Конструктор КСП",
            "📊 Анализ и визуализация (EDA)",
            "🤖 ML-Прогноз уровня ученика",
            "📷 AI-Проверка по фото",
            "👤 Генератор характеристик",
            "⚡ Разминки и интерактивы",
        ],
        "no_key": "Пожалуйста, введите ваш рабочий Gemini API Key в боковом меню слева!",
        "ai_lang_prompt": "Напиши ответ строго на русском языке.",
        
        "source": "Источник данных:",
        "source_options": ["Google Таблица (ссылка)", "Excel-файл (.xlsx)"],
        "sheet_link": "Ссылка на Google Таблицу:",
        "upload_excel": "Загрузите Excel-файл (листы: Банк_вопросов, Ученики):",
        "success_excel": "Excel-файл успешно прочитан!",
        "settings": "Параметры генерации",
        "easy": "Легких вопросов:",
        "med": "Средних вопросов:",
        "hard": "Сложных вопросов:",
        "gen_word": "Сгенерировать варианты в Word",
        "wait_ai": "ИИ обрабатывает банк вопросов и формирует варианты...",
        "student_lbl": "Ученик(ца):",
        "task_lbl": "Задание",
        "answer_lbl": "Ответ: ____________________",
        "keys_title": "КЛЮЧИ (ДЛЯ УЧИТЕЛЯ)",
        "done": "Документы успешно созданы!",
        "download_cards": "Скачать Карточки (Word)",
        "download_keys": "Скачать Ключи (Word)",
        
        "subject": "Учебный предмет:",
        "grade": "Класс:",
        "quarters": "Количество четвертей:",
        "hours": "Часов в неделю:",
        "total_lessons": "Всего академических часов:",
        "source_pdf_text": "Источник тем:",
        "pdf_opt": ["Загрузить PDF-файл", "Ввести темы текстом"],
        "topics_lbl": "Перечень тем:",
        "gen_ktp": "Сгенерировать КТП в Word",
        "wait_ktp": "ИИ анализирует материалы и формирует КТП...",
        "download_ktp": "Скачать КТП (Word)",
        
        "teacher_name": "ФИО преподавателя:",
        "topic_lbl": "Тема урока:",
        "target_lbl": "Цели обучения (ЦО):",
        "gen_ksp": "Сгенерировать КСП в Word",
        "wait_ksp": "ИИ методист разрабатывает структуру урока...",
        "download_ksp": "Скачать КСП (Word)",
        
        "eda_title": "Анализ и визуализация успеваемости класса",
        "eda_sub": "Загрузите файл с оценками для построения описательной статистики и графиков",
        "eda_load": "Загрузить датасет (.xlsx)",
        "eda_select": "Показатель для анализа:",
        "eda_btn": "Построить аналитику",
        "eda_wait": "Расчет метрик и построение графиков...",
        "hist": "Гистограмма плотности распределения",
        "box": "Ящик с усами (Boxplot анализа)",
        
        "ml_title": "Интеллектуальный прогноз успеваемости",
        "ml_sub": "Классификация уровня учащихся с помощью алгоритма Random Forest",
        "ml_txt": "Введите ключевые показатели студента:",
        "att": "Посещаемость (%):",
        "hw": "Выполнение ДЗ (%):",
        "test": "Средний балл тестов:",
        "activity": "Академическая активность:",
        "act_opts": ["Низкая", "Средняя", "Высокая"],
        "ml_btn": "Выполнить ML-прогноз",
        "ml_wait": "Модель классификации анализирует данные...",
        "rec": "Прогноз модели:",
        
        "photo_title": "Мультимодальная проверка работ",
        "photo_load": "Загрузите фото письменной работы:",
        "photo_check": "Запустить проверку",
        "photo_wait": "ИИ распознает рукописный текст и ищет ошибки...",
        
        "char_title": "Генератор педагогических характеристик",
        "char_sub": "Автоматизированное составление отчета на основе параметров",
        "name_lbl": "ФИО учащегося:",
        "cls_lbl": "Академическая группа / Класс:",
        "att_lbl": "Посещаемость (%):",
        "perf_lbl": "Успеваемость:",
        "perf_opts": ["Отличник", "Ударник", "Занимается средне", "Имеет академические задолженности"],
        "beh_lbl": "Дисциплина:",
        "beh_opts": ["Дисциплинирован, примерное поведение", "Спокойный, исполнительный", "Иногда нарушает дисциплину", "Требует повышенного педагогического внимания"],
        "traits_lbl": "Дополнительные достижения и качества:",
        "char_btn": "Сгенерировать характеристику",
        "char_wait": "Формирование текста характеристики...",
        
        "warm_title": "AI-Генератор разминок (Icebreakers)",
        "warm_sub": "Интерактивные упражнения для вовлечения аудитории",
        "warm_top": "Тема занятия:",
        "warm_time": "Тайминг (минут):",
        "warm_btn": "Подобрать активности",
        "warm_wait": "Генерация интерактивных заданий...",
    },
    "kk": {
        "page_title": "AI Мұғалім Көмекшісі",
        "sidebar_title": "🎓 EduAI Platform",
        "api_subheader": "🔑 ЖИ қолжетімділік кілті",
        "api_help": "Барлық құралдар үшін кілтті бір рет енгізіңіз",
        "api_expander": "ℹ️ API кілтін қалай алуға болады?",
        "tools_subheader": "🛠️ Жүйе модульдері",
        "footer": "✨ Оқытушылар үшін әзірленген",
        "menu": [
            "📝 Тапсырма карточкаларын жасау",
            "📅 КТП AI-Генераторы",
            "📋 ҚМЖ (КСП) AI-Конструкторы",
            "📊 Талдау және визуализация (EDA)",
            "🤖 Оқушы деңгейін ML болжау",
            "📷 Фото арқылы AI тексеру",
            "👤 Мінездеме генераторы",
            "⚡ Сергіту сәттері мен интерактив",
        ],
        "no_key": "Сол жақ мәзірге жарамды Gemini API Key енгізіңіз!",
        "ai_lang_prompt": "Жауапты қатаң түрде қазақ тілінде жаз.",
        
        "source": "Дереккөз:",
        "source_options": ["Google Кесте (сілтеме)", "Excel-файл (.xlsx)"],
        "sheet_link": "Google кестенің сілтемесі:",
        "upload_excel": "Excel файлын жүктеңіз (парақтар: Банк_вопросов, Ученики):",
        "success_excel": "Excel файлы сәтті оқылды!",
        "settings": "Генерация параметрлері",
        "easy": "Жеңіл сұрақтар:",
        "med": "Орташа сұрақтар:",
        "hard": "Қиын сұрақтар:",
        "gen_word": "Word форматында нұсқалар жасау",
        "wait_ai": "ЖИ сұрақтар банкін өңдеп жатыр...",
        "student_lbl": "Оқушы:",
        "task_lbl": "Тапсырма",
        "answer_lbl": "Жауап: ____________________",
        "keys_title": "ЖАУАПТАР (МҰҒАЛІМГЕ)",
        "done": "Құжаттар сәтті дайындалды!",
        "download_cards": "Карточкаларды жүктеу (Word)",
        "download_keys": "Жауаптарды жүктеу (Word)",
        
        "subject": "Оқу пәні:",
        "grade": "Сынып / Курс:",
        "quarters": "Тоқсан саны:",
        "hours": "Аптасына сағат:",
        "total_lessons": "Барлық академиялық сағат:",
        "source_pdf_text": "Тақырыптар көзі:",
        "pdf_opt": ["PDF файлын жүктеу", "Тақырыптарды мәтінмен енгізу"],
        "topics_lbl": "Тақырыптар тізімі:",
        "gen_ktp": "Word форматында КТП құру",
        "wait_ktp": "ЖИ материалдарды талдап, КТП жасауда...",
        "download_ktp": "КТП жүктеу (Word)",
        
        "teacher_name": "Оқытушының А.Т.Ә.:",
        "topic_lbl": "Сабақ тақырыбы:",
        "target_lbl": "Оқыту мақсаттары (ОМ):",
        "gen_ksp": "Word форматында ҚМЖ құру",
        "wait_ksp": "ЖИ әдіскер сабақ жоспарын әзірлеуде...",
        "download_ksp": "ҚМЖ жүктеу (Word)",
        
        "eda_title": "Сынып үлгерімін талдау және визуализация",
        "eda_sub": "Статистика мен графиктер құру үшін бағалар файлын жүктеңіз",
        "eda_load": "Датасетті жүктеу (.xlsx)",
        "eda_select": "Талдау көрсеткіші:",
        "eda_btn": "Аналитика құру",
        "eda_wait": "Метрикалар есептеліп, графиктер салынуда...",
        "hist": "Бөлініс тығыздығының гистограммасы",
        "box": "Жәшік диаграммасы (Boxplot)",
        
        "ml_title": "Зияткерлік үлгерім болжамы",
        "ml_sub": "Random Forest алгоритмі арқылы оқушы деңгейін жіктеу",
        "ml_txt": "Оқушының негізгі көрсеткіштерін енгізіңіз:",
        "att": "Сабаққа қатысу (%):",
        "hw": "Үй жұмысы (%):",
        "test": "Тесттердің орташа балы:",
        "activity": "Академиялық белсенділік:",
        "act_opts": ["Төмен", "Орташа", "Жоғары"],
        "ml_btn": "ML-болжам жасау",
        "ml_wait": "Жіктеу моделі деректерді өңдеуде...",
        "rec": "Модель болжамы:",
        
        "photo_title": "Жұмыстарды мультимодальды тексеру",
        "photo_load": "Жазбаша жұмыс фотосы:",
        "photo_check": "Тексеруді бастау",
        "photo_wait": "Қолжазба танылып, қателер тексерілуде...",
        
        "char_title": "Педагогикалық мінездеме генераторы",
        "char_sub": "Көрсеткіштер негізінде автоматты түрде мінездеме құру",
        "name_lbl": "Оқушының А.Т.Ә.:",
        "cls_lbl": "Академиялық топ / Сынып:",
        "att_lbl": "Қатысу (%):",
        "perf_lbl": "Үлгерім:",
        "perf_opts": ["Үздік", "Екпінді", "Орташа оқиды", "Академиялық қарыздары бар"],
        "beh_lbl": "Тәртіп:",
        "beh_opts": ["Тәртіпті, үлгілі мінез-құлық", "Сабырлы, жауапкершілікті", "Кейде тәртіпті бұзады", "Ерекше педагогикалық бақылауды қажет етеді"],
        "traits_lbl": "Қосымша жетістіктер мен қасиеттер:",
        "char_btn": "Мінездеме құру",
        "char_wait": "Мінездеме мәтіні дайындалуда...",
        
        "warm_title": "AI Сергіту сәттері мен Icebreakers",
        "warm_sub": "Аудиторияны қызықтыруға арналған интерактивті жаттығулар",
        "warm_top": "Сабақ тақырыбы:",
        "warm_time": "Тайминг (минут):",
        "warm_btn": "Белсенділіктерді таңдау",
        "warm_wait": "Интерактивті тапсырмалар жасалуда...",
    }
}

# ==========================================
# 1. СОВРЕМЕННЫЙ ДИЗАЙН (UI / CSS)
# ==========================================
st.set_page_config(page_title="AI-Помощник Учителя", page_icon="🎓", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', sans-serif;
    }
    
    /* Общий фон приложения */
    .stApp {
        background: linear-gradient(135deg, #f4f6f9 0%, #edf2f7 100%);
    }
    
    /* Основной контейнер с эффектом карточки */
    .block-container {
        background-color: #ffffff;
        border-radius: 24px;
        padding: 3rem 3.5rem !important;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.04);
        border: 1px solid rgba(226, 232, 240, 0.8);
        margin-top: 2rem;
        margin-bottom: 2rem;
    }
    
    /* Боковая панель */
    [data-testid="stSidebar"] {
        background-color: #0f172a;
        color: #ffffff;
        border-right: none;
    }
    
    [data-testid="stSidebar"] .stMarkdown h1, 
    [data-testid="stSidebar"] .stMarkdown h2, 
    [data-testid="stSidebar"] .stMarkdown h3, 
    [data-testid="stSidebar"] label {
        color: #e2e8f0 !important;
    }
    
    /* Кнопки */
    .stButton>button {
        border-radius: 14px;
        background: linear-gradient(135deg, #6366f1 0%, #4f46e5 100%);
        color: white !important;
        font-weight: 600;
        padding: 0.6rem 1.2rem;
        border: none;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(99, 102, 241, 0.25);
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(99, 102, 241, 0.4);
        background: linear-gradient(135deg, #4f46e5 100%, #4338ca 100%);
    }
    
    /* Заголовки */
    h1 {
        color: #0f172a;
        font-weight: 800;
        letter-spacing: -0.5px;
    }
    
    h3, h4 {
        color: #334155;
        font-weight: 700;
    }
    
    /* Поля ввода */
    .stTextInput>div>div>input, .stNumberInput>div>div>input, .stSelectbox>div>div>div {
        border-radius: 12px;
        border: 1px solid #cbd5e1;
    }
    
    /* Скрываем лишние элементы Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

DEFAULT_API_KEY = ""

# ==========================================
# 2. БОКОВОЕ МЕНЮ И ЯЗЫК
# ==========================================
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/1972/1972413.png", width=70)

lang_choice = st.sidebar.selectbox("🌐 Тіл / Язык интерфейса:", ["Русский", "Қазақша"], index=0)
lang = "ru" if lang_choice == "Русский" else "kk"
t = translations[lang]

st.sidebar.markdown(f"### {t['sidebar_title']}")
st.sidebar.markdown("---")

st.sidebar.subheader(t["api_subheader"])
global_api_key = st.sidebar.text_input("Gemini API Key:", type="password", help=t["api_help"])

with st.sidebar.expander(t["api_expander"]):
    st.markdown("""
    1. Зайдите на [Google AI Studio](https://aistudio.google.com/app/apikey).
    2. Войдите через Google-аккаунт.
    3. Нажмите кнопку **Create API key**.
    4. Скопируйте и вставьте ключ выше.
    """)

active_key = global_api_key if global_api_key else DEFAULT_API_KEY

st.sidebar.markdown("---")
st.sidebar.subheader(t["tools_subheader"])

menu_choice = st.sidebar.radio(
    "Navigation:",
    t["menu"],
    label_visibility="collapsed"
)

st.sidebar.markdown("---")
st.sidebar.caption(t["footer"])


# ==========================================
# МОДУЛЬ 1: ГЕНЕРАТОР КАРТОЧЕК
# ==========================================
if menu_choice in ["📝 Генератор карточек", "📝 Тапсырма карточкаларын жасау"]:
    st.title(f"📝 {menu_choice}")
    st.markdown(f"#### {'Автоматическая генерация индивидуальных вариантов в Word' if lang=='ru' else 'Word форматында жеке нұсқаларды автоматты түрде жасау'}")
    st.divider()

    source_type = st.radio(t["source"], t["source_options"], horizontal=True)
    df_questions, df_students = None, None

    if "Google" in source_type:
        default_url = "https://docs.google.com/spreadsheets/d/1fJKlRP7YY3r6DFjd_PuLXFIKkg3GdSRAM9Rxwq502e8/edit?usp=sharing"
        sheet_url = st.text_input(f"🔗 {t['sheet_link']}", value=default_url)
        def extract_sheet_id(url):
            try: return url.split("/d/")[1].split("/")[0] if "/d/" in url else url
            except: return None
        if sheet_url:
            sheet_id = extract_sheet_id(sheet_url)
            if sheet_id:
                try:
                    def get_url(name):
                        enc = urllib.parse.quote(name)
                        return f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={enc}"
                    df_questions = pd.read_csv(get_url("Банк_вопросов"))
                    df_students = pd.read_csv(get_url("Ученики"))
                except: pass
    else:
        uploaded_excel = st.file_uploader(f"📂 {t['upload_excel']}", type=["xlsx"])
        if uploaded_excel:
            try:
                xls = pd.ExcelFile(uploaded_excel)
                df_questions = pd.read_excel(xls, 'Банк_вопросов')
                df_students = pd.read_excel(xls, 'Ученики')
                st.success(t["success_excel"])
            except Exception as e: st.error(f"Ошибка: {e}")

    if df_questions is not None and df_students is not None:
        st.markdown(f"#### ⚙️ {t['settings']}")
        col1, col2, col3 = st.columns(3)
        with col1: count_easy = st.number_input(f"🟢 {t['easy']}", min_value=0, max_value=5, value=1)
        with col2: count_med = st.number_input(f"🟡 {t['med']}", min_value=0, max_value=5, value=1)
        with col3: count_hard = st.number_input(f"🔴 {t['hard']}", min_value=0, max_value=5, value=1)

        if st.button(f"🚀 {t['gen_word']}", type="primary", use_container_width=True):
            if not active_key: st.warning(t["no_key"])
            else:
                try:
                    with st.spinner(f"⏳ {t['wait_ai']}"):
                        df_questions.columns = df_questions.columns.astype(str).str.strip().str.lower()
                        df_students.columns = df_students.columns.astype(str).str.strip().str.lower()
                        rename_dict = {}
                        for col in df_questions.columns:
                            if "сложн" in col or "күрдел" in col: rename_dict[col] = "сложность"
                            elif "тем" in col or "тақырып" in col: rename_dict[col] = "тема"
                            elif "вопрос" in col or "сұрақ" in col: rename_dict[col] = "вопрос"
                            elif "ответ" in col or "жауап" in col: rename_dict[col] = "ответ"
                        df_questions = df_questions.rename(columns=rename_dict)
                        student_col = [c for c in df_students.columns if "фио" in c or "аты" in c]
                        student_col_name = student_col[0] if student_col else df_students.columns[0]
                        students = df_students[student_col_name].dropna().tolist()

                    doc_students = Document()
                    structure = {"Легкий": count_easy, "Средний": count_med, "Сложный": count_hard}
                    teacher_keys = []
                    for student in students:
                        variant_questions = []
                        for level, count in structure.items():
                            if count > 0:
                                subset = df_questions[df_questions["сложность"].astype(str).str.strip().str.capitalize() == level]
                                if len(subset) == 0: subset = df_questions
                                variant_questions.append(subset.sample(n=min(count, len(subset))))
                        student_variant = pd.concat(variant_questions).reset_index(drop=True)
                        title = doc_students.add_heading("Проверочная работа" if lang=="ru" else "Бақылау жұмысы", level=2)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_info = doc_students.add_paragraph()
                        p_info.add_run(f"{t['student_lbl']} {student}").bold = True
                        for idx, row in student_variant.iterrows():
                            p_q = doc_students.add_paragraph()
                            p_q.add_run(f"{t['task_lbl']} {idx + 1} ").bold = True
                            p_q.add_run(f"{row['вопрос']}\n")
                            p_q.add_run(t["answer_lbl"])
                            teacher_keys.append({"Ученик": student, "№ Задания": idx + 1, "Ответ": row["ответ"]})
                        doc_students.add_paragraph("--------------------------------------------------")

                    bio_students = io.BytesIO()
                    doc_students.save(bio_students)
                    bio_students.seek(0)
                    doc_teacher = Document()
                    doc_teacher.add_heading(t["keys_title"], level=1)
                    df_keys = pd.DataFrame(teacher_keys)
                    curr_st = ""
                    for _, row in df_keys.iterrows():
                        if row["Ученик"] != curr_st:
                            curr_st = row["Ученик"]
                            doc_teacher.add_paragraph().add_run(f"\n👤 {row['Ученик']}").bold = True
                        doc_teacher.add_paragraph(f"  • {t['task_lbl']} {row['№ Задания']}: {row['Ответ']}")
                    bio_teacher = io.BytesIO()
                    doc_teacher.save(bio_teacher)
                    bio_teacher.seek(0)

                    st.success(f"🎉 {t['done']}")
                    col_d1, col_d2 = st.columns(2)
                    with col_d1: st.download_button(f"📄 {t['download_cards']}", bio_students, "Карточки.docx", use_container_width=True)
                    with col_d2: st.download_button(f"🔑 {t['download_keys']}", bio_teacher, "Ключи.docx", use_container_width=True)
                except Exception as e: st.error(f"Ошибка: {e}")

# ==========================================
# МОДУЛЬ 2: AI-ГЕНЕРАТОР КТП (WORD)
# ==========================================
elif menu_choice in ["📅 AI-Генератор КТП", "📅 КТП AI-Генераторы"]:
    st.title(f"📅 {menu_choice}")
    st.divider()
    col_p1, col_p2 = st.columns(2)
    with col_p1:
        subject = st.text_input(t["subject"], "Информатика")
        grade = st.number_input(t["grade"], 1, 11, 8)
    with col_p2:
        quarters_count = st.selectbox(t["quarters"], [1, 2, 3, 4], index=0)
        hours_per_week = st.number_input(t["hours"], 1, 5, 2)
    quarters_weeks = {q: st.number_input(f"{q}-я четверть (недель):" if lang=="ru" else f"{q}-ші тоқсан (апта):", 1, 15, 8) for q in range(1, quarters_count + 1)}
    total_all_lessons = sum(q_w * hours_per_week for q_w in quarters_weeks.values())
    st.info(f"💡 {t['total_lessons']} **{total_all_lessons}**")

    source_type = st.radio(t["source_pdf_text"], t["pdf_opt"], horizontal=True)
    uploaded_pdf, textbook_content = None, ""
    if "PDF" in source_type: uploaded_pdf = st.file_uploader("📂 PDF:", type=["pdf"])
    else: textbook_content = st.text_area(t["topics_lbl"], "1. Инфо 2. Алгоритмы", height=100)

    if st.button(f"🚀 {t['gen_ktp']}", type="primary", use_container_width=True):
        if not active_key: st.warning(t["no_key"])
        else:
            try:
                with st.spinner(f"⏳ {t['wait_ktp']}"):
                    genai.configure(api_key=active_key)
                    model = genai.GenerativeModel("gemini-3.6-flash")
                    prompt = f"{t['ai_lang_prompt']} Составь КТП по предмету {subject}, {grade} класс, уроков: {total_all_lessons}. Темы: {textbook_content}. Верни строго JSON массив: [{{'quarter':1, 'lesson_num':1, 'topic':'', 'targets':'', 'homework':''}}]"
                    response = model.generate_content([prompt, uploaded_pdf] if uploaded_pdf else prompt)
                    match = re.search(r'\[.*\]', response.text, re.DOTALL)
                    ktp_data = json.loads(match.group(0) if match else response.text)
                    
                    df_ktp = pd.DataFrame(ktp_data)
                    st.dataframe(df_ktp, use_container_width=True)

                    doc_ktp = Document()
                    section = doc_ktp.sections[-1]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width

                    title = doc_ktp.add_paragraph()
                    r = title.add_run("КАЛЕНДАРНО-ТЕМАТИЧЕСКОЕ ПЛАНИРОВАНИЕ (КТП)\n" if lang=="ru" else "КҮНТІЗБЕЛІК-ТАҚЫРЫПТЫҚ ЖОСПАР (КТП)\n")
                    r.font.bold = True
                    r.font.size = Pt(14)
                    r.font.name = 'Times New Roman'
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc_ktp.add_paragraph()

                    table = doc_ktp.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    headers = ["Четверть / Тоқсан", "№", "Тема урока / Сабақ тақырыбы", "Цели обучения / Оқыту мақсаттары", "ДЗ / Үй жұмысы"]
                    for i, h in enumerate(headers):
                        table.rows[0].cells[i].text = h

                    for item in ktp_data:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(item.get("quarter", ""))
                        row_cells[1].text = str(item.get("lesson_num", ""))
                        row_cells[2].text = str(item.get("topic", ""))
                        row_cells[3].text = str(item.get("targets", ""))
                        row_cells[4].text = str(item.get("homework", ""))

                    bio = io.BytesIO()
                    doc_ktp.save(bio)
                    bio.seek(0)
                    
                    st.success(f"🎉 {t['done']}")
                    st.download_button(f"📄 {t['download_ktp']}", bio, f"КТП_{subject}.docx", use_container_width=True)
            except Exception as e: st.error(f"Ошибка: {e}")

# ==========================================
# МОДУЛЬ 3: AI-КОНСТРУКТОР КСП
# ==========================================
elif menu_choice in ["📋 AI-Конструктор КСП", "📋 ҚМЖ (КСП) AI-Конструкторы"]:
    st.title(f"📋 {menu_choice}")
    st.divider()
    col_k1, col_k2 = st.columns(2)
    with col_k1:
        teacher_name = st.text_input(t["teacher_name"], "Иванов И.И.")
        subject_ksp = st.text_input(t["subject"], "Информатика")
        grade_ksp = st.number_input(t["grade"], 1, 11, 8)
    with col_k2:
        topic_ksp = st.text_input(t["topic_lbl"], "Двоичная система")
        target_ksp = st.text_input(t["target_lbl"], "8.1.1.1 Перевод чисел")

    if st.button(f"🚀 {t['gen_ksp']}", type="primary", use_container_width=True):
        if not active_key: st.warning(t["no_key"])
        else:
            try:
                with st.spinner(f"⏳ {t['wait_ksp']}"):
                    genai.configure(api_key=active_key)
                    model = genai.GenerativeModel("gemini-3.6-flash")
                    prompt = f"{t['ai_lang_prompt']} Создай план урока по предмету {subject_ksp}, тема {topic_ksp}. Верни строго JSON: {{\"lesson_targets\":\"...\", \"eval_criteria\":\"...\", \"stages\":[{{\"time\":\"Начало\", \"teacher\":\"...\", \"student\":\"...\", \"eval\":\"...\", \"resources\":\"...\"}}]}}"
                    response = model.generate_content(prompt)
                    match = re.search(r'\{.*\}', response.text, re.DOTALL)
                    ksp_data = json.loads(match.group(0) if match else response.text)

                    doc_ksp = Document()
                    section = doc_ksp.sections[-1]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width
                    
                    title = doc_ksp.add_paragraph()
                    r = title.add_run("КРАТКОСРОЧНЫЙ ПЛАН УРОКА (ҚМЖ)")
                    r.font.bold = True
                    r.font.size = Pt(14)
                    r.font.name = 'Times New Roman'
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    t_table = doc_ksp.add_table(rows=7, cols=2)
                    t_table.style = 'Table Grid'
                    info = [("Мұғалім / Учитель:", teacher_name), ("Пән / Предмет:", subject_ksp), ("Сынып / Класс:", str(grade_ksp)), ("Тақырып / Тема:", topic_ksp), ("ОМ / ЦО:", target_ksp), ("Мақсат / Цели:", ksp_data.get("lesson_targets","")), ("Критерий:", ksp_data.get("eval_criteria",""))]
                    for idx, (l, v) in enumerate(info):
                        t_table.rows[idx].cells[0].text = l
                        t_table.rows[idx].cells[1].text = str(v)

                    doc_ksp.add_paragraph()
                    s_table = doc_ksp.add_table(rows=1, cols=5)
                    s_table.style = 'Table Grid'
                    headers = ["Этап / Кезең", "Действия учителя", "Действия ученика", "Оценивание", "Ресурсы"]
                    for i, h in enumerate(headers): s_table.rows[0].cells[i].text = h

                    for stg in ksp_data.get("stages", []):
                        row = s_table.add_row().cells
                        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = stg.get("time",""), stg.get("teacher",""), stg.get("student",""), stg.get("eval",""), stg.get("resources","")

                    bio = io.BytesIO()
                    doc_ksp.save(bio)
                    bio.seek(0)
                    st.success(f"🎉 {t['done']}")
                    st.download_button(f"📄 {t['download_ksp']}", bio, f"ҚМЖ_{topic_ksp}.docx", use_container_width=True)
            except Exception as e: st.error(f"Ошибка: {e}")

# ==========================================
# МОДУЛЬ 4: EDA
# ==========================================
elif menu_choice in ["📊 Анализ и визуализация (EDA)", "📊 Талдау және визуализация (EDA)"]:
    st.title(f"📊 {menu_choice}")
    st.markdown(f"#### {t['eda_sub']}")
    st.divider()
    uploaded_eda = st.file_uploader(f"📂 {t['eda_load']}", type=["xlsx"])
    if uploaded_eda:
        df_eda = pd.read_excel(uploaded_eda)
        st.write(df_eda.head())
        num_cols = df_eda.select_dtypes(include=['number']).columns.tolist()
        if num_cols:
            col = st.selectbox(t["eda_select"], num_cols)
            if st.button(f"📈 {t['eda_btn']}", type="primary"):
                with st.spinner(f"⏳ {t['eda_wait']}"):
                    st.write(df_eda[col].describe())
                    col_g1, col_g2 = st.columns(2)
                    with col_g1:
                        st.markdown(f"##### {t['hist']}")
                        fig, ax = plt.subplots(figsize=(6, 4))
                        sns.histplot(df_eda[col], kde=True, ax=ax, color='#6366f1')
                        st.pyplot(fig)
                    with col_g2:
                        st.markdown(f"##### {t['box']}")
                        fig, ax = plt.subplots(figsize=(6, 4))
                        sns.boxplot(y=df_eda[col], ax=ax, color='#a5b4fc')
                        st.pyplot(fig)

# ==========================================
# МОДУЛЬ 5: ML
# ==========================================
elif menu_choice in ["🤖 ML-Прогноз уровня ученика", "🤖 Оқушы деңгейін ML болжау"]:
    st.title(f"🤖 {menu_choice}")
    st.markdown(f"#### {t['ml_sub']}")
    st.divider()
    st.write(t["ml_txt"])
    col_m1, col_m2 = st.columns(2)
    with col_m1:
        att = st.slider(t["att"], 50, 100, 85)
        hw = st.slider(t["hw"], 0, 100, 75)
    with col_m2:
        test = st.slider(t["test"], 0, 100, 80)
        activity = st.selectbox(t["activity"], t["act_opts"])
        act_val = 1 if activity in ["Низкая", "Төмен"] else (2 if activity in ["Средняя", "Орташа"] else 3)

    if st.button(f"🔮 {t['ml_btn']}", type="primary", use_container_width=True):
        with st.spinner(f"⏳ {t['ml_wait']}"):
            X_train = [[60, 50, 55, 1], [90, 85, 88, 3], [70, 60, 65, 2], [95, 95, 92, 3]]
            y_train = ["Группа поддержки" if lang=="ru" else "Қолдау тобы", "Продвинутый" if lang=="ru" else "Жоғары", "Стандартный" if lang=="ru" else "Стандартты", "Продвинутый" if lang=="ru" else "Жоғары"]
            model = RandomForestClassifier(random_state=42).fit(X_train, y_train)
            pred = model.predict([[att, hw, test, act_val]])[0]
        st.success(f"🎯 {t['rec']} **{pred}**")

# ==========================================
# МОДУЛЬ 6: AI-ПРОВЕРКА ПО ФОТО
# ==========================================
elif menu_choice in ["📷 AI-Проверка по фото", "📷 Фото арқылы AI тексеру"]:
    st.title(f"📷 {menu_choice}")
    st.divider()
    img = st.file_uploader(f"📂 {t['photo_load']}", type=["jpg", "png"])
    if img and st.button(t["photo_check"], type="primary"):
        if not active_key: st.warning(t["no_key"])
        else:
            with st.spinner(f"⏳ {t['photo_wait']}"):
                genai.configure(api_key=active_key)
                model = genai.GenerativeModel("gemini-3.6-flash")
                res = model.generate_content([f"{t['ai_lang_prompt']} Проверь работу ученика:", Image.open(img)])
                st.markdown(res.text)

# ==========================================
# МОДУЛЬ 7: ХАРАКТЕРИСТИКА
# ==========================================
elif menu_choice in ["👤 Генератор характеристик", "👤 Мінездеме генераторы"]:
    st.title(f"👤 {menu_choice}")
    st.markdown(f"#### {t['char_sub']}")
    st.divider()
    col_h1, col_h2 = st.columns(2)
    with col_h1:
        name = st.text_input(t["name_lbl"], "Иванов Иван")
        cls = st.text_input(t["cls_lbl"], "8 «А»")
        att = st.slider(t["att_lbl"], 0, 100, 90)
    with col_h2:
        perf = st.selectbox(t["perf_lbl"], t["perf_opts"])
        beh = st.selectbox(t["beh_lbl"], t["beh_opts"])

    traits = st.text_area(t["traits_lbl"], "...")

    if st.button(f"🚀 {t['char_btn']}", type="primary", use_container_width=True):
        if not active_key: st.warning(t["no_key"])
        else:
            with st.spinner(f"⏳ {t['char_wait']}"):
                genai.configure(api_key=active_key)
                model = genai.GenerativeModel("gemini-3.6-flash")
                prompt = f"{t['ai_lang_prompt']} Напиши официальную характеристику на ученика {name}, класс {cls}. Посещаемость: {att}%, успеваемость: {perf}, поведение: {beh}, доп: {traits}."
                res = model.generate_content(prompt)
                st.markdown(res.text)

# ==========================================
# МОДУЛЬ 8: РАЗМИНКИ
# ==========================================
elif menu_choice in ["⚡ Разминки и интерактивы", "⚡ Сергіту сәттері мен интерактив"]:
    st.title(f"⚡ {menu_choice}")
    st.markdown(f"#### {t['warm_sub']}")
    st.divider()
    col_w1, col_w2 = st.columns(2)
    with col_w1:
        top = st.text_input(t["warm_top"], "Алгоритмы")
    with col_w2:
        tm = st.slider(t["warm_time"], 2, 10, 5)

    if st.button(f"🚀 {t['warm_btn']}", type="primary", use_container_width=True):
        if not active_key: st.warning(t["no_key"])
        else:
            with st.spinner(f"⏳ {t['warm_wait']}"):
                genai.configure(api_key=active_key)
                model = genai.GenerativeModel("gemini-3.6-flash")
                prompt = f"{t['ai_lang_prompt']} Предложи 3 разминки на тему {top} на {tm} минут."
                res = model.generate_content(prompt)
                st.markdown(res.text)

# ==========================================
# 📊 GOOGLE ANALYTICS СЧЕТЧИК
# ==========================================
GA_TRACKING_ID = "G-0EW4TYEDKE" 
ga_component = f"""
<!DOCTYPE html>
<html>
<head>
    <script async src="https://www.googletagmanager.com/gtag/js?id={GA_TRACKING_ID}"></script>
    <script>
      window.dataLayer = window.dataLayer || [];
      function gtag(){{dataLayer.push(arguments);}}
      gtag('js', new Date());
      gtag('config', '{GA_TRACKING_ID}');
    </script>
</head>
<body>
</body>
</html>
"""
components.html(ga_component, height=0, width=0)
